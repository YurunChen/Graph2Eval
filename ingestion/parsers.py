"""
Document parsers for different file formats
"""

import io
import re
from abc import ABC, abstractmethod
from dataclasses import dataclass
from typing import List, Dict, Any, Optional, Union
from pathlib import Path

import PyPDF2
import pdfplumber
from bs4 import BeautifulSoup
from PIL import Image
import docx
from loguru import logger


@dataclass
class ParsedElement:
    """Represents a parsed document element"""
    element_type: str  # paragraph, table, heading, figure, metadata
    content: str
    metadata: Dict[str, Any]
    page_num: Optional[int] = None
    bbox: Optional[tuple] = None  # (x0, y0, x1, y1)
    parent_id: Optional[str] = None
    element_id: Optional[str] = None


@dataclass 
class DocumentStructure:
    """Complete document structure after parsing"""
    elements: List[ParsedElement]
    metadata: Dict[str, Any]
    total_pages: int
    file_path: str


class DocumentParser(ABC):
    """Abstract base class for document parsers"""
    
    @abstractmethod
    def parse(self, file_path: Union[str, Path]) -> DocumentStructure:
        """Parse document and return structured representation"""
        pass
    
    @abstractmethod
    def extract_text(self, file_path: Union[str, Path]) -> str:
        """Extract plain text from document"""
        pass


class PDFParser(DocumentParser):
    """Enhanced PDF parser with table and figure detection"""
    
    def __init__(self, extract_tables: bool = True, extract_images: bool = True, 
                 image_config: Optional[Dict[str, Any]] = None, output_dir: Optional[Path] = None):
        self.extract_tables = extract_tables
        self.extract_images = extract_images
        self.image_config = image_config or {}
        self.output_dir = output_dir  # Override default output directory
        self.current_pdf_path = None  # Store current PDF path for image extraction
        
    def parse(self, file_path: Union[str, Path]) -> DocumentStructure:
        """Parse PDF with structure detection"""
        file_path = Path(file_path)
        logger.info(f"Parsing PDF: {file_path}")
        
        # Store PDF path for image extraction
        self.current_pdf_path = str(file_path)
        
        elements = []
        metadata = {}
        
        try:
            with pdfplumber.open(file_path) as pdf:
                metadata = {
                    "title": pdf.metadata.get("Title", ""),
                    "author": pdf.metadata.get("Author", ""),
                    "creator": pdf.metadata.get("Creator", ""),
                    "creation_date": pdf.metadata.get("CreationDate", ""),
                    "total_pages": len(pdf.pages)
                }
                
                for page_num, page in enumerate(pdf.pages, 1):
                    # Extract text blocks
                    text_elements = self._extract_text_elements(page, page_num)
                    elements.extend(text_elements)
                    
                    # Extract tables
                    if self.extract_tables:
                        table_elements = self._extract_tables(page, page_num)
                        elements.extend(table_elements)
                    
                    # Extract images
                    if self.extract_images:
                        image_elements = self._extract_images(page, page_num)
                        elements.extend(image_elements)
                        
        except Exception as e:
            logger.error(f"Error parsing PDF {file_path}: {e}")
            raise
            
        return DocumentStructure(
            elements=elements,
            metadata=metadata,
            total_pages=metadata.get("total_pages", 0),
            file_path=str(file_path)
        )
    
    def _extract_text_elements(self, page, page_num: int) -> List[ParsedElement]:
        """Extract text elements with hierarchy detection"""
        elements = []
        
        # Get text with bounding boxes
        text_objects = page.extract_words(keep_blank_chars=True)
        
        if not text_objects:
            return elements
            
        # Group words into lines and paragraphs
        lines = self._group_words_to_lines(text_objects)
        paragraphs = self._group_lines_to_paragraphs(lines)
        
        for i, para in enumerate(paragraphs):
            text = " ".join([word["text"] for word in para])
            
            # Determine element type based on formatting
            element_type = self._classify_text_element(para, text)
            
            # Calculate bounding box
            bbox = self._calculate_bbox(para)
            
            element = ParsedElement(
                element_type=element_type,
                content=text.strip(),
                metadata={
                    "font_size": para[0].get("size", 0) if para else 0,
                    "font": para[0].get("fontname", "") if para else "",
                    "word_count": len(text.split()),
                    "char_count": len(text)
                },
                page_num=page_num,
                bbox=bbox,
                element_id=f"page_{page_num}_text_{i}"
            )
            
            if element.content.strip():  # Only add non-empty elements
                elements.append(element)
                
        return elements
    
    def _extract_tables(self, page, page_num: int) -> List[ParsedElement]:
        """Extract tables from page"""
        elements = []
        
        try:
            tables = page.extract_tables()
            
            for i, table in enumerate(tables):
                if not table:
                    continue
                    
                # Convert table to string representation
                table_text = self._table_to_text(table)
                
                # Extract table caption from surrounding text
                caption = self._extract_table_caption(page, table, i)
                
                # Combine caption with table content
                if caption:
                    full_content = f"Table {i+1}: {caption}\n\n{table_text}"
                else:
                    full_content = table_text
                
                # Get table bbox (approximate)
                bbox = page.bbox  # Placeholder - could be improved
                
                element = ParsedElement(
                    element_type="table",
                    content=full_content,
                    metadata={
                        "rows": len(table),
                        "cols": len(table[0]) if table else 0,
                        "table_data": table,
                        "caption": caption,
                        "table_number": i + 1
                    },
                    page_num=page_num,
                    bbox=bbox,
                    element_id=f"page_{page_num}_table_{i}"
                )
                
                elements.append(element)
                
        except Exception as e:
            logger.warning(f"Error extracting tables from page {page_num}: {e}")
            
        return elements
    
    def _extract_images(self, page, page_num: int) -> List[ParsedElement]:
        """Extract images from page and save to files"""
        elements = []
        
        try:
            # Use pymupdf for better image extraction
            import fitz
            
            # Open the PDF file directly with pymupdf
            if not self.current_pdf_path:
                raise ValueError("PDF path not available for image extraction")
            doc = fitz.open(self.current_pdf_path)
            page_index = page_num - 1  # Convert to 0-based index
            
            # Get the pymupdf page
            pymupdf_page = doc[page_index]
            
            # Get images using pymupdf
            image_list = pymupdf_page.get_images()
            
            for i, img_info in enumerate(image_list):
                image_path = None
                
                logger.debug(f"Image {i} on page {page_num}: {img_info}")
                
                # Save image if configured
                if self.image_config.get('save_images', False):
                    image_path = self._save_image_with_pymupdf(doc, img_info, page_num, i)
                
                # Extract image caption from surrounding text
                caption = self._extract_image_caption(page, i)
                
                # Create content with caption
                if caption:
                    content = f"Figure {i+1}: {caption}"
                else:
                    # Generate a more descriptive placeholder when no caption is found
                    content = f"Figure {i+1}: Image from page {page_num} (no caption available)"
                
                element = ParsedElement(
                    element_type="figure",
                    content=content,
                    metadata={
                        "width": img_info[2] if img_info else 0,
                        "height": img_info[3] if img_info else 0,
                        "image_path": image_path,
                        "caption": caption,
                        "figure_number": i + 1
                    },
                    page_num=page_num,
                    bbox=(0, 0, img_info[2] if img_info else 0, img_info[3] if img_info else 0),
                    element_id=f"page_{page_num}_image_{i}"
                )
                
                elements.append(element)
            
            doc.close()
                
        except ImportError:
            logger.warning("pymupdf not available, falling back to pdfplumber image extraction")
            # Fallback to pdfplumber
            elements = self._extract_images_fallback(page, page_num)
        except Exception as e:
            logger.warning(f"Error extracting images from page {page_num}: {e}")
            
        return elements
    
    def _extract_images_fallback(self, page, page_num: int) -> List[ParsedElement]:
        """Fallback image extraction using pdfplumber"""
        elements = []
        
        try:
            images = page.images
            
            for i, img in enumerate(images):
                image_path = None
                
                # Check if this is a real image or just an image mask
                is_image_mask = img.get('imagemask', False)
                if is_image_mask:
                    logger.info(f"Skipping image mask on page {page_num}, image {i}")
                    continue
                
                # Save image if configured
                if self.image_config.get('save_images', False):
                    image_path = self._save_image(img, page_num, i)
                
                element = ParsedElement(
                    element_type="figure",
                    content=f"[Image: {img.get('name', 'unnamed')}]",
                    metadata={
                        "image_data": img,
                        "width": img.get("width", 0),
                        "height": img.get("height", 0),
                        "image_path": image_path
                    },
                    page_num=page_num,
                    bbox=(img.get("x0", 0), img.get("top", 0), 
                          img.get("x1", 0), img.get("bottom", 0)),
                    element_id=f"page_{page_num}_image_{i}"
                )
                
                elements.append(element)
                
        except Exception as e:
            logger.warning(f"Error in fallback image extraction: {e}")
            
        return elements
    
    def _save_image(self, img, page_num: int, image_index: int) -> Optional[str]:
        """Save image to file and return the path"""
        try:
            # Get image configuration
            if self.output_dir:
                # Use benchmark output directory
                output_dir = self.output_dir
            else:
                # Use default configuration
                output_dir = Path(self.image_config.get('image_output_dir', 'data/images'))
            
            image_format = self.image_config.get('image_format', 'png')
            max_size = self.image_config.get('max_image_size', 2048)
            compress = self.image_config.get('compress_images', True)
            quality = self.image_config.get('image_quality', 85)
            
            # Create output directory
            output_dir.mkdir(parents=True, exist_ok=True)
            
            # Generate filename
            filename = f"page_{page_num}_image_{image_index}.{image_format}"
            image_path = output_dir / filename
            
            # Get image data
            image_stream = img.get('stream')
            if not image_stream:
                logger.warning(f"No image data found for image {image_index} on page {page_num}")
                return None
            
            # Convert PDFStream to bytes
            try:
                if hasattr(image_stream, 'get_data'):
                    # PDFStream object
                    image_data = image_stream.get_data()
                elif hasattr(image_stream, 'read'):
                    # File-like object
                    image_stream.seek(0)  # Reset to beginning
                    image_data = image_stream.read()
                elif isinstance(image_stream, bytes):
                    # Already bytes
                    image_data = image_stream
                else:
                    # Try to convert to bytes
                    image_data = bytes(image_stream)
                
                # Validate that we have actual image data
                if not image_data or len(image_data) < 100:
                    logger.warning(f"Image data too small or empty: {len(image_data) if image_data else 0} bytes")
                    return None
                    
                # Debug: log the first few bytes to understand the format
                logger.debug(f"Image data first 20 bytes: {image_data[:20]}")
                
                # Check if it's valid image data by looking for common image headers
                if (image_data.startswith(b'\x89PNG\r\n\x1a\n') or  # PNG
                    image_data.startswith(b'\xff\xd8\xff') or      # JPEG
                    image_data.startswith(b'GIF8') or             # GIF
                    image_data.startswith(b'BM') or               # BMP
                    image_data.startswith(b'%!PS') or             # PostScript
                    image_data.startswith(b'%PDF')):              # PDF
                    logger.info(f"Valid image format detected")
                else:
                    # Try to convert PDF image data to PNG using PIL
                    try:
                        from PIL import Image
                        import io
                        
                        # Try to open as image (PIL can handle many formats)
                        img = Image.open(io.BytesIO(image_data))
                        
                        # Convert to PNG and save
                        img_buffer = io.BytesIO()
                        img.save(img_buffer, format='PNG')
                        image_data = img_buffer.getvalue()
                        
                        logger.info(f"Successfully converted image to PNG format")
                        
                    except Exception as pil_error:
                        logger.warning(f"Failed to convert image with PIL: {pil_error}")
                        logger.warning(f"Image data doesn't have valid image header and conversion failed")
                        return None
                    
            except Exception as e:
                logger.warning(f"Failed to extract image data: {e}")
                return None
            
            # Save image
            with open(image_path, 'wb') as f:
                f.write(image_data)
            
            logger.info(f"Saved image: {image_path}")
            return str(image_path)
            
        except Exception as e:
            logger.error(f"Error saving image: {e}")
            return None
    
    def _save_image_with_pymupdf(self, doc, img_info, page_num: int, image_index: int) -> Optional[str]:
        """Save image using pymupdf and return the path"""
        try:
            # Get image configuration
            if self.output_dir:
                # Use benchmark output directory
                output_dir = self.output_dir
            else:
                # Use default configuration
                output_dir = Path(self.image_config.get('image_output_dir', 'data/images'))
            image_format = self.image_config.get('image_format', 'png')
            
            # Create output directory
            output_dir.mkdir(parents=True, exist_ok=True)
            
            # Generate filename
            filename = f"page_{page_num}_image_{image_index}.{image_format}"
            image_path = output_dir / filename
            
            # Extract image using pymupdf
            import fitz
            
            xref = img_info[0]
            pix = fitz.Pixmap(doc, xref)
            
            if pix.n - pix.alpha < 4:  # GRAY or RGB
                img_data = pix.tobytes("png")
            else:  # CMYK: convert to RGB first
                pix1 = fitz.Pixmap(fitz.csRGB, pix)
                img_data = pix1.tobytes("png")
                pix1 = None
            
            # Save image
            with open(image_path, 'wb') as f:
                f.write(img_data)
            
            pix = None
            
            logger.info(f"Saved image with pymupdf: {image_path}")
            return str(image_path)
            
        except Exception as e:
            logger.error(f"Error saving image with pymupdf: {e}")
            return None
    
    def _group_words_to_lines(self, words: List[Dict]) -> List[List[Dict]]:
        """Group words into lines based on y-coordinate"""
        if not words:
            return []
            
        # Sort by y-coordinate (top to bottom)
        words_sorted = sorted(words, key=lambda w: (w["top"], w["x0"]))
        
        lines = []
        current_line = [words_sorted[0]]
        
        for word in words_sorted[1:]:
            # Check if word is on the same line (similar y-coordinate)
            if abs(word["top"] - current_line[-1]["top"]) < 5:  # 5 pixel tolerance
                current_line.append(word)
            else:
                lines.append(current_line)
                current_line = [word]
                
        if current_line:
            lines.append(current_line)
            
        return lines
    
    def _group_lines_to_paragraphs(self, lines: List[List[Dict]]) -> List[List[Dict]]:
        """Group lines into paragraphs based on spacing"""
        if not lines:
            return []
            
        paragraphs = []
        current_para = lines[0]
        
        for line in lines[1:]:
            # Calculate spacing between lines
            prev_bottom = max(word["bottom"] for word in current_para)
            curr_top = min(word["top"] for word in line)
            spacing = curr_top - prev_bottom
            
            # If spacing is large, start new paragraph
            if spacing > 10:  # 10 pixel threshold
                paragraphs.append(current_para)
                current_para = line
            else:
                current_para.extend(line)
                
        if current_para:
            paragraphs.append(current_para)
            
        return paragraphs
    
    def _classify_text_element(self, words: List[Dict], text: str) -> str:
        """Classify text element type based on formatting and content"""
        if not words:
            return "paragraph"
            
        avg_font_size = sum(w.get("size", 12) for w in words) / len(words)
        
        # Check for headings
        if (avg_font_size > 14 or 
            any([re.match(r'^\d+\.?\s+[A-Z]', text), 
                re.match(r'^[A-Z][A-Z\s]+$', text),
                re.match(r'^\d+\.\d+', text)])):
            return "heading"
            
        # Check for list items
        if re.match(r'^\s*[•\-\*]\s+', text) or re.match(r'^\s*\d+\.\s+', text):
            return "list_item"
            
        return "paragraph"
    
    def _calculate_bbox(self, words: List[Dict]) -> tuple:
        """Calculate bounding box for group of words"""
        if not words:
            return (0, 0, 0, 0)
            
        x0 = min(w["x0"] for w in words)
        y0 = min(w["top"] for w in words)
        x1 = max(w["x1"] for w in words)
        y1 = max(w["bottom"] for w in words)
        
        return (x0, y0, x1, y1)
    
    def _table_to_text(self, table: List[List[str]]) -> str:
        """Convert table data to text representation"""
        if not table:
            return ""
            
        # Create markdown-style table
        lines = []
        for i, row in enumerate(table):
            if row:  # Skip empty rows
                clean_row = [str(cell).strip() if cell else "" for cell in row]
                lines.append(" | ".join(clean_row))
                
                # Add header separator
                if i == 0:
                    separator = " | ".join(["---"] * len(clean_row))
                    lines.append(separator)
                    
        return "\n".join(lines)
    
    def _extract_table_caption(self, page, table, table_index: int) -> str:
        """Extract table caption from surrounding text"""
        try:
            # Get all text from the page
            page_text = page.extract_text()
            if not page_text:
                return ""
            
            # Look for common table caption patterns
            caption_patterns = [
                r'Table\s+\d+[:\-]?\s*(.+?)(?=\n\n|\nTable|\nFigure|\n[A-Z]|\Z)',
                r'Table\s+\d+\.\d+[:\-]?\s*(.+?)(?=\n\n|\nTable|\nFigure|\n[A-Z]|\Z)',
                r'Table\s+[A-Z]\.\d+[:\-]?\s*(.+?)(?=\n\n|\nTable|\nFigure|\n[A-Z]|\Z)',
                r'^\s*Table\s+\d+[:\-]?\s*(.+?)$',
            ]
            
            for pattern in caption_patterns:
                matches = re.findall(pattern, page_text, re.IGNORECASE | re.MULTILINE | re.DOTALL)
                if matches:
                    # Return the first match that seems like a caption
                    for match in matches:
                        caption = match.strip()
                        if len(caption) > 10 and len(caption) < 200:  # Reasonable caption length
                            return caption
            
            # If no pattern matches, look for text near the table
            # This is a simplified approach - could be improved with better positioning
            return ""
            
        except Exception as e:
            logger.debug(f"Error extracting table caption: {e}")
            return ""
    
    def _extract_image_caption(self, page, image_index: int) -> str:
        """Extract image caption from surrounding text"""
        try:
            # Get all text from the page
            page_text = page.extract_text()
            if not page_text:
                return ""
            
            # Look for common figure caption patterns
            caption_patterns = [
                r'Figure\s+\d+[:\-]?\s*(.+?)(?=\n\n|\nTable|\nFigure|\n[A-Z]|\Z)',
                r'Figure\s+\d+\.\d+[:\-]?\s*(.+?)(?=\n\n|\nTable|\nFigure|\n[A-Z]|\Z)',
                r'Figure\s+[A-Z]\.\d+[:\-]?\s*(.+?)(?=\n\n|\nTable|\nFigure|\n[A-Z]|\Z)',
                r'Fig\.\s+\d+[:\-]?\s*(.+?)(?=\n\n|\nTable|\nFigure|\n[A-Z]|\Z)',
                r'^\s*Figure\s+\d+[:\-]?\s*(.+?)$',
            ]
            
            for pattern in caption_patterns:
                matches = re.findall(pattern, page_text, re.IGNORECASE | re.MULTILINE | re.DOTALL)
                if matches:
                    # Return the first match that seems like a caption
                    for match in matches:
                        caption = match.strip()
                        if len(caption) > 10 and len(caption) < 200:  # Reasonable caption length
                            return caption
            
            # If no pattern matches, look for text near the image
            # This is a simplified approach - could be improved with better positioning
            return ""
            
        except Exception as e:
            logger.debug(f"Error extracting image caption: {e}")
            return ""
    
    def extract_text(self, file_path: Union[str, Path]) -> str:
        """Extract plain text from PDF"""
        file_path = Path(file_path)
        
        try:
            with pdfplumber.open(file_path) as pdf:
                text_parts = []
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
                        
                return "\n\n".join(text_parts)
                
        except Exception as e:
            logger.error(f"Error extracting text from PDF {file_path}: {e}")
            return ""


class HTMLParser(DocumentParser):
    """HTML parser with semantic structure extraction"""
    
    def __init__(self, extract_links: bool = True, extract_images: bool = True,
                 image_config: Optional[Dict[str, Any]] = None, output_dir: Optional[Path] = None):
        self.extract_links = extract_links
        self.extract_images = extract_images
        self.image_config = image_config or {}
        self.output_dir = output_dir  # Override default output directory
        
    def parse(self, file_path: Union[str, Path]) -> DocumentStructure:
        """Parse HTML with semantic structure"""
        file_path = Path(file_path)
        
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
            
        soup = BeautifulSoup(content, 'html.parser')
        
        elements = []
        metadata = self._extract_metadata(soup)
        
        # Extract structured elements
        elements.extend(self._extract_headings(soup))
        elements.extend(self._extract_paragraphs(soup))
        elements.extend(self._extract_lists(soup))
        elements.extend(self._extract_tables(soup))
        
        if self.extract_links:
            elements.extend(self._extract_links(soup))
            
        if self.extract_images:
            elements.extend(self._extract_images(soup))
            
        return DocumentStructure(
            elements=elements,
            metadata=metadata,
            total_pages=1,  # HTML is single "page"
            file_path=str(file_path)
        )
    
    def _extract_metadata(self, soup: BeautifulSoup) -> Dict[str, Any]:
        """Extract HTML metadata"""
        metadata = {}
        
        # Title
        title = soup.find('title')
        if title:
            metadata['title'] = title.get_text().strip()
            
        # Meta tags
        meta_tags = soup.find_all('meta')
        for meta in meta_tags:
            name = meta.get('name') or meta.get('property')
            content = meta.get('content')
            if name and content:
                metadata[name] = content
                
        return metadata
    
    def _extract_headings(self, soup: BeautifulSoup) -> List[ParsedElement]:
        """Extract heading elements"""
        elements = []
        
        for i, heading in enumerate(soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6'])):
            element = ParsedElement(
                element_type="heading",
                content=heading.get_text().strip(),
                metadata={
                    "tag": heading.name,
                    "level": int(heading.name[1]),
                    "id": heading.get('id', ''),
                    "class": heading.get('class', [])
                },
                element_id=f"heading_{i}"
            )
            elements.append(element)
            
        return elements
    
    def _extract_paragraphs(self, soup: BeautifulSoup) -> List[ParsedElement]:
        """Extract paragraph elements"""
        elements = []
        
        for i, para in enumerate(soup.find_all('p')):
            text = para.get_text().strip()
            if text:
                element = ParsedElement(
                    element_type="paragraph",
                    content=text,
                    metadata={
                        "tag": "p",
                        "class": para.get('class', []),
                        "word_count": len(text.split())
                    },
                    element_id=f"paragraph_{i}"
                )
                elements.append(element)
                
        return elements
    
    def _extract_lists(self, soup: BeautifulSoup) -> List[ParsedElement]:
        """Extract list elements"""
        elements = []
        
        for i, lst in enumerate(soup.find_all(['ul', 'ol'])):
            items = [li.get_text().strip() for li in lst.find_all('li')]
            content = '\n'.join(f"• {item}" for item in items if item)
            
            if content:
                element = ParsedElement(
                    element_type="list",
                    content=content,
                    metadata={
                        "tag": lst.name,
                        "item_count": len(items),
                        "class": lst.get('class', [])
                    },
                    element_id=f"list_{i}"
                )
                elements.append(element)
                
        return elements
    
    def _extract_tables(self, soup: BeautifulSoup) -> List[ParsedElement]:
        """Extract table elements with captions"""
        elements = []
        
        for i, table in enumerate(soup.find_all('table')):
            rows = []
            for tr in table.find_all('tr'):
                cells = [td.get_text().strip() for td in tr.find_all(['td', 'th'])]
                if cells:
                    rows.append(cells)
                    
            if rows:
                # Extract table caption
                caption = self._extract_html_table_caption(table)
                
                # Convert to markdown-style table
                table_content = self._table_to_text(rows)
                
                # Combine caption with table content
                if caption:
                    content = f"Table {i+1}: {caption}\n\n{table_content}"
                else:
                    content = table_content
                
                element = ParsedElement(
                    element_type="table",
                    content=content,
                    metadata={
                        "tag": "table",
                        "rows": len(rows),
                        "cols": len(rows[0]) if rows else 0,
                        "class": table.get('class', []),
                        "table_data": rows,
                        "caption": caption,
                        "table_number": i + 1
                    },
                    element_id=f"table_{i}"
                )
                elements.append(element)
                
        return elements
    
    def _extract_html_table_caption(self, table) -> str:
        """Extract table caption from HTML table element"""
        try:
            # Look for <caption> tag
            caption_tag = table.find('caption')
            if caption_tag:
                return caption_tag.get_text().strip()
            
            # Look for preceding elements that might be captions
            # Check previous sibling elements
            prev_element = table.find_previous_sibling()
            if prev_element:
                text = prev_element.get_text().strip()
                # Check if it looks like a caption
                if (text and len(text) < 200 and 
                    any(keyword in text.lower() for keyword in ['table', 'data', 'results', 'summary'])):
                    return text
            
            # Look for following elements (some captions come after)
            next_element = table.find_next_sibling()
            if next_element:
                text = next_element.get_text().strip()
                if (text and len(text) < 200 and 
                    any(keyword in text.lower() for keyword in ['table', 'data', 'results', 'summary'])):
                    return text
            
            return ""
            
        except Exception as e:
            logger.debug(f"Error extracting HTML table caption: {e}")
            return ""
    
    def _extract_links(self, soup: BeautifulSoup) -> List[ParsedElement]:
        """Extract link elements"""
        elements = []
        
        for i, link in enumerate(soup.find_all('a', href=True)):
            text = link.get_text().strip()
            href = link['href']
            
            if text and href:
                element = ParsedElement(
                    element_type="link",
                    content=f"{text} ({href})",
                    metadata={
                        "tag": "a",
                        "href": href,
                        "text": text,
                        "class": link.get('class', [])
                    },
                    element_id=f"link_{i}"
                )
                elements.append(element)
                
        return elements
    
    def _extract_images(self, soup: BeautifulSoup) -> List[ParsedElement]:
        """Extract image elements and save to files"""
        elements = []
        
        for i, img in enumerate(soup.find_all('img')):
            src = img.get('src', '')
            alt = img.get('alt', '')
            image_path = None
            
            # Save image if configured and src is a local file
            if self.image_config.get('save_images', False) and src:
                image_path = self._save_html_image(src, i)
            
            # Extract image caption
            caption = self._extract_html_image_caption(img)
            
            # Create content with caption
            if caption:
                content = f"Figure {i+1}: {caption}\n[Image: {alt or src}]"
            else:
                content = f"[Image: {alt or src}]"
            
            element = ParsedElement(
                element_type="figure",
                content=content,
                metadata={
                    "tag": "img",
                    "src": src,
                    "alt": alt,
                    "width": img.get('width', ''),
                    "height": img.get('height', ''),
                    "class": img.get('class', []),
                    "image_path": image_path,
                    "caption": caption,
                    "figure_number": i + 1
                },
                element_id=f"image_{i}"
            )
            elements.append(element)
            
        return elements
    
    def _extract_html_image_caption(self, img) -> str:
        """Extract image caption from HTML img element"""
        try:
            # First check if alt text is descriptive enough to be a caption
            alt_text = img.get('alt', '').strip()
            if alt_text and len(alt_text) > 10 and len(alt_text) < 200:
                return alt_text
            
            # Look for title attribute
            title = img.get('title', '').strip()
            if title and len(title) > 10 and len(title) < 200:
                return title
            
            # Look for preceding elements that might be captions
            prev_element = img.find_previous_sibling()
            if prev_element:
                text = prev_element.get_text().strip()
                if (text and len(text) < 200 and 
                    any(keyword in text.lower() for keyword in ['figure', 'image', 'photo', 'chart', 'graph'])):
                    return text
            
            # Look for following elements (some captions come after)
            next_element = img.find_next_sibling()
            if next_element:
                text = next_element.get_text().strip()
                if (text and len(text) < 200 and 
                    any(keyword in text.lower() for keyword in ['figure', 'image', 'photo', 'chart', 'graph'])):
                    return text
            
            # Look for parent element that might contain caption
            parent = img.parent
            if parent:
                # Check if parent has a caption-like class or id
                parent_class = ' '.join(parent.get('class', []))
                parent_id = parent.get('id', '')
                if any(keyword in (parent_class + ' ' + parent_id).lower() for keyword in ['caption', 'figure', 'image']):
                    text = parent.get_text().strip()
                    if text and len(text) < 200:
                        return text
            
            return ""
            
        except Exception as e:
            logger.debug(f"Error extracting HTML image caption: {e}")
            return ""
    
    def _save_html_image(self, src: str, image_index: int) -> Optional[str]:
        """Save HTML image to file and return the path"""
        try:
            # Get image configuration
            if self.output_dir:
                # Use benchmark output directory
                output_dir = self.output_dir
            else:
                # Use default configuration
                output_dir = Path(self.image_config.get('image_output_dir', 'data/images'))
            image_format = self.image_config.get('image_format', 'png')
            
            # Create output directory
            output_dir.mkdir(parents=True, exist_ok=True)
            
            # Generate filename
            filename = f"html_image_{image_index}.{image_format}"
            image_path = output_dir / filename
            
            # Handle different src types
            if src.startswith('data:'):
                # Data URL
                import base64
                header, data = src.split(',', 1)
                image_data = base64.b64decode(data)
            elif src.startswith('http'):
                # Remote URL - skip for now (could add download functionality)
                logger.info(f"Skipping remote image: {src}")
                return None
            else:
                # Local file path
                src_path = Path(src)
                if src_path.exists():
                    with open(src_path, 'rb') as f:
                        image_data = f.read()
                else:
                    logger.warning(f"Image file not found: {src}")
                    return None
            
            # Save image
            with open(image_path, 'wb') as f:
                f.write(image_data)
            
            logger.info(f"Saved HTML image: {image_path}")
            return str(image_path)
            
        except Exception as e:
            logger.error(f"Error saving HTML image: {e}")
            return None
    
    def _table_to_text(self, rows: List[List[str]]) -> str:
        """Convert table rows to text representation"""
        if not rows:
            return ""
            
        lines = []
        for i, row in enumerate(rows):
            lines.append(" | ".join(row))
            if i == 0:  # Add header separator
                separator = " | ".join(["---"] * len(row))
                lines.append(separator)
                
        return "\n".join(lines)
    
    def extract_text(self, file_path: Union[str, Path]) -> str:
        """Extract plain text from HTML"""
        file_path = Path(file_path)
        
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
            
        soup = BeautifulSoup(content, 'html.parser')
        
        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.decompose()
            
        return soup.get_text()


class DOCXParser(DocumentParser):
    """DOCX parser for Word documents"""
    
    def parse(self, file_path: Union[str, Path]) -> DocumentStructure:
        """Parse DOCX document"""
        file_path = Path(file_path)
        
        doc = docx.Document(file_path)
        elements = []
        
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if text:
                # Determine element type based on style
                element_type = "heading" if paragraph.style.name.startswith('Heading') else "paragraph"
                
                element = ParsedElement(
                    element_type=element_type,
                    content=text,
                    metadata={
                        "style": paragraph.style.name,
                        "level": self._get_heading_level(paragraph.style.name) if element_type == "heading" else 0
                    },
                    element_id=f"paragraph_{i}"
                )
                elements.append(element)
                
        # Extract tables
        for i, table in enumerate(doc.tables):
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):  # Skip empty rows
                    rows.append(cells)
                    
            if rows:
                # Extract table caption (simplified for DOCX)
                caption = self._extract_docx_table_caption(table, i)
                
                table_content = self._table_to_text(rows)
                
                # Combine caption with table content
                if caption:
                    content = f"Table {i+1}: {caption}\n\n{table_content}"
                else:
                    content = table_content
                
                element = ParsedElement(
                    element_type="table",
                    content=content,
                    metadata={
                        "rows": len(rows),
                        "cols": len(rows[0]) if rows else 0,
                        "table_data": rows,
                        "caption": caption,
                        "table_number": i + 1
                    },
                    element_id=f"table_{i}"
                )
                elements.append(element)
        
        metadata = {
            "title": doc.core_properties.title or "",
            "author": doc.core_properties.author or "",
            "created": str(doc.core_properties.created) if doc.core_properties.created else "",
            "modified": str(doc.core_properties.modified) if doc.core_properties.modified else ""
        }
        
        return DocumentStructure(
            elements=elements,
            metadata=metadata,
            total_pages=1,
            file_path=str(file_path)
        )
    
    def _get_heading_level(self, style_name: str) -> int:
        """Extract heading level from style name"""
        if 'Heading' in style_name:
            try:
                return int(style_name.split()[-1])
            except (ValueError, IndexError):
                return 1
        return 0
    
    def _extract_docx_table_caption(self, table, table_index: int) -> str:
        """Extract table caption from DOCX table (simplified)"""
        try:
            # For DOCX, we can look at the table's title property if available
            # This is a simplified implementation - DOCX caption extraction is complex
            if hasattr(table, 'title') and table.title:
                return table.title.strip()
            
            # Look for table caption in the table's first row if it looks like a caption
            if table.rows:
                first_row = table.rows[0]
                if len(first_row.cells) == 1:  # Single cell might be a caption
                    cell_text = first_row.cells[0].text.strip()
                    if (cell_text and len(cell_text) < 200 and 
                        any(keyword in cell_text.lower() for keyword in ['table', 'data', 'results', 'summary'])):
                        return cell_text
            
            return ""
            
        except Exception as e:
            logger.debug(f"Error extracting DOCX table caption: {e}")
            return ""
    
    def _table_to_text(self, rows: List[List[str]]) -> str:
        """Convert table rows to text representation"""
        if not rows:
            return ""
            
        lines = []
        for i, row in enumerate(rows):
            lines.append(" | ".join(row))
            if i == 0:  # Add header separator
                separator = " | ".join(["---"] * len(row))
                lines.append(separator)
                
        return "\n".join(lines)
    
    def extract_text(self, file_path: Union[str, Path]) -> str:
        """Extract plain text from DOCX"""
        file_path = Path(file_path)
        doc = docx.Document(file_path)
        
        text_parts = []
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                text_parts.append(text)
                
        return "\n\n".join(text_parts)
